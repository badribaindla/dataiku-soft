from io import BytesIO

from docx.table import _Cell

from dataiku.base.utils import safe_unicode_str
from .document_handler import DocumentHandler
from dataiku.doctor.docgen.extractor.docx_parser import DocxParser
from docx.enum.style import WD_STYLE_TYPE
import logging
from docx.oxml.shared import OxmlElement, qn
import json

class Text(object):
    def __init__(self, text):
        self.inner_text = text

    def apply(self, paragraph, document, initial_run, table_style = None):
        """
        Insert the placeholder data as a text inside the location defined by the initial_run inside the paragraph.
        :param paragraph: the location of the placeholder
        :param document: the current document, it will be used to retrieve style (unused)
        :param initial_run: the run of the placeholder
        :param table_style: a potential style to apply to the current table. (unused)
        :return: a new paragraph after the text insertion.
        """
        if initial_run:
            run = paragraph.add_run(self.inner_text, initial_run.style)
            DocumentHandler.copy_font(initial_run, run)
        else:
            paragraph.add_run(self.inner_text)
        return paragraph


class Table(object):
    def __init__(self, cells):
        self.cells = cells

    @staticmethod
    def move_table_after(table, paragraph):
        tbl, p = table._tbl, paragraph._p
        p.addnext(tbl)

    def apply(self, paragraph, document, initial_run, table_style = None):
        """
        Insert the placeholder data as a table inside the location defined by the initial_run inside the paragraph.
        :param paragraph: the location of the placeholder
        :param document: the current document, it will be used to retrieve style
        :param initial_run: the run of the placeholder
        :param table_style: a potential style to apply to the current table.
        :return: a new paragraph after the table insertion.
        """
        output_paragraph = DocumentHandler.insert_paragraph_after(paragraph, style=paragraph.style)

        if len(self.cells) < 1:
            table = document.add_table(rows=1, cols=1)
        else:
            max_col_size = 1
            for r in range(0, len(self.cells)):
                if len(self.cells[r]) > max_col_size:
                    max_col_size = len(self.cells[r])
            table = document.add_table(rows=len(self.cells), cols=max_col_size)

        if table_style is not None:
            table.style = table_style.style.name
            self.__copy_table_style__(table, table_style)
        elif hasattr(initial_run, "style") and initial_run.style and initial_run.style.type == WD_STYLE_TYPE.TABLE:
            table.style = initial_run.style

        # we have to split the current paragraph in two, to place the table in the middle, in its own format
        self.move_table_after(table, paragraph)
        for r in range(0, len(self.cells)):
            table_cells = table.rows[r].cells
            for c in range(0, len(self.cells[r])):
                table_cells[c].text = self.cells[r][c]
        return output_paragraph

    def __copy_table_style__(self, table, table_style):
        """
        Duplicate the style from table_style into table
        :param table: a empty table that need a style
        :param table_style: a reference tab with a style
        """

        # Start by copying the global style of the table
        self.__copy_table_metadata__(table, table_style)

        # Report existing columns/rows
        row_id = 0

        while row_id < len(table_style.rows) and row_id < len(table.rows):
            col_id = 0
            table_style_cells = table_style.rows[row_id].cells
            table_cells = table.rows[row_id].cells
            while col_id < len(table_style_cells) and col_id < len(table_cells):
                self.__copy_cell_style__(table_style_cells[col_id], table_cells[col_id])
                col_id += 1
            row_id += 1

        # new rows
        row_id = len(table_style.rows)
        if row_id > 0:
            while row_id < len(table.rows):
                col_id = 0
                previous_table_cells = table.rows[row_id - 1].cells
                table_cells = table.rows[row_id].cells
                while col_id < len(table_cells):
                    self.__copy_cell_style__(previous_table_cells[col_id], table_cells[col_id])
                    col_id += 1
                row_id += 1

        # new columns
        row_id = 0
        while row_id < len(table.rows):
            col_id = len(table_style.rows[0].cells)
            table_cells = table.rows[row_id].cells
            if col_id > 0:
                while col_id < len(table_cells):
                    self.__copy_cell_style__(table_cells[col_id - 1], table_cells[col_id])
                    col_id += 1
            row_id += 1

    def __extract_attribute__(self, reference_tc_pr, output_element, path, attribute_name, default=None):
        """
        Use Xpath to extract an element from reference_tc_pr and duplicate it on output_element
        :param reference_tc_pr: input element ad a tc_pr with attributes
        :param output_element: empty output element
        :param path: XML path
        :param attribute_name: name of the attribute
        :param default: default value if no value were found
        """
        attribute = reference_tc_pr.xpath(path + "/@" + attribute_name)
        if len(attribute) > 0:
            output_element.set(qn(attribute_name), attribute[0])
        elif default:
            output_element.set(qn(attribute_name), default)

    def __copy_table_metadata__(self, table, table_style):
        """
        Copy the table metadata, like "is the first row a header"
        :param table:
        :param table_style:
        :return:
        """
        # useful for debug, display the input xml: logging.debug("XML from _tbl %s", table_style.rows._tbl.xml)
        from_tbl_pr = table_style.rows._tbl.tblPr
        if len(from_tbl_pr.xpath("w:tblLook")) > 0:
            to_tbl_pr = table.rows._tbl.tblPr
            if len(to_tbl_pr.xpath("w:tblLook")) > 0:
                to_tbl_pr.remove_all("w:tblLook")
            w_tbl_look = OxmlElement("w:tblLook")

            self.__extract_attribute__(from_tbl_pr, w_tbl_look, "w:tblLook", "w:val")
            self.__extract_attribute__(from_tbl_pr, w_tbl_look, "w:tblLook", "w:firstRow")
            self.__extract_attribute__(from_tbl_pr, w_tbl_look, "w:tblLook", "w:lastRow")
            self.__extract_attribute__(from_tbl_pr, w_tbl_look, "w:tblLook", "w:firstColumn")
            self.__extract_attribute__(from_tbl_pr, w_tbl_look, "w:tblLook", "w:lastColumn")
            self.__extract_attribute__(from_tbl_pr, w_tbl_look, "w:tblLook", "w:noHBand")
            self.__extract_attribute__(from_tbl_pr, w_tbl_look, "w:tblLook", "w:noVBand")
            to_tbl_pr.append(w_tbl_look)


    def __copy_border__(self, reference_tc_pr, element_name):
        """
        Copy the different element of a border into a new OxmlElement
        :param reference_tc_pr: element we want to get border information from
        :param element_name: the type of border we want to get (w:top, w:bottom, w:left or w:right)
        :return: a new OxmlElement that copy the style of the input
        """
        element = OxmlElement(element_name)

        tc_borders_path = "w:tcBorders/"
        self.__extract_attribute__(reference_tc_pr, element, tc_borders_path + element_name, "w:val")
        self.__extract_attribute__(reference_tc_pr, element, tc_borders_path + element_name, "w:color")
        self.__extract_attribute__(reference_tc_pr, element, tc_borders_path + element_name, "w:sz")
        self.__extract_attribute__(reference_tc_pr, element, tc_borders_path + element_name, "w:space")
        self.__extract_attribute__(reference_tc_pr, element, tc_borders_path + element_name, "w:shadow")
        self.__extract_attribute__(reference_tc_pr, element, tc_borders_path + element_name, "w:frame")
        return element

    def __copy_margin__(self, reference_tc_pr, element_name):
        """
        Copy the different element of a margin into a new OxmlElement
        :param reference_tc_pr: element we want to get border information from
        :param element_name: the type of border we want to get (w:top, w:bottom, w:left or w:right)
        :return: a new OxmlElement that copy the style of the input
        """
        element = OxmlElement(element_name)

        tc_mar_path = "w:tcMar/"
        self.__extract_attribute__(reference_tc_pr, element, tc_mar_path + element_name, "w:fill")
        self.__extract_attribute__(reference_tc_pr, element, tc_mar_path + element_name, "w:w")
        return element

    def __copy_cell_style__(self, from_cell, to_cell):
        """
        Copy the style of "from_cell" into the "to_cell"
        :param from_cell: a cell containing a pretty style
        :param to_cell: a cell without style
        """
        from_tc_pr = from_cell._tc.get_or_add_tcPr()
        to_tc_pr = to_cell._tc.get_or_add_tcPr()
        # useful for debug, display the input xml: logging.debug("XML tc_pr %s", from_tc_pr.xml)

        # Copy the configuration (has header/footer/...)
        if len(from_tc_pr.xpath("w:cnfStyle")) > 0:
            w_cnf_style = OxmlElement("w:cnfStyle")
            self.__extract_attribute__(from_tc_pr, w_cnf_style, "w:cnfStyle", "w:val")
            self.__extract_attribute__(from_tc_pr, w_cnf_style, "w:cnfStyle", "w:firstRow")
            self.__extract_attribute__(from_tc_pr, w_cnf_style, "w:cnfStyle", "w:lastRow")
            self.__extract_attribute__(from_tc_pr, w_cnf_style, "w:cnfStyle", "w:firstColumn")
            self.__extract_attribute__(from_tc_pr, w_cnf_style, "w:cnfStyle", "w:lastColumn")
            self.__extract_attribute__(from_tc_pr, w_cnf_style, "w:cnfStyle", "w:oddVBand")
            self.__extract_attribute__(from_tc_pr, w_cnf_style, "w:cnfStyle", "w:evenVBand")
            self.__extract_attribute__(from_tc_pr, w_cnf_style, "w:cnfStyle", "w:oddHBand")
            self.__extract_attribute__(from_tc_pr, w_cnf_style, "w:cnfStyle", "w:evenHBand")
            self.__extract_attribute__(from_tc_pr, w_cnf_style, "w:cnfStyle", "w:firstRowFirstColumn")
            self.__extract_attribute__(from_tc_pr, w_cnf_style, "w:cnfStyle", "w:firstRowLastColumn")
            self.__extract_attribute__(from_tc_pr, w_cnf_style, "w:cnfStyle", "w:lastRowFirstColumn")
            self.__extract_attribute__(from_tc_pr, w_cnf_style, "w:cnfStyle", "w:lastRowLastColumn")
            to_tc_pr.append(w_cnf_style)


        # Copy the color of the text and the background
        if len(from_tc_pr.xpath("w:shd")) > 0:
            w_shd = OxmlElement("w:shd")
            self.__extract_attribute__(from_tc_pr, w_shd, "w:shd", "w:fill", "auto")
            self.__extract_attribute__(from_tc_pr, w_shd, "w:shd", "w:color", "auto")
            self.__extract_attribute__(from_tc_pr, w_shd, "w:shd", "w:themeColor")
            self.__extract_attribute__(from_tc_pr, w_shd, "w:shd", "w:themeFill")
            self.__extract_attribute__(from_tc_pr, w_shd, "w:shd", "w:themeShade")
            self.__extract_attribute__(from_tc_pr, w_shd, "w:shd", "w:themeFillShade")
            to_tc_pr.append(w_shd)

        # Copy the borders
        if len(from_tc_pr.xpath("w:tcBorders")) > 0:
            w_tc_borders = OxmlElement("w:tcBorders")
            w_tc_borders.append(self.__copy_border__(from_tc_pr, "w:top"))
            w_tc_borders.append(self.__copy_border__(from_tc_pr, "w:bottom"))
            w_tc_borders.append(self.__copy_border__(from_tc_pr, "w:left"))
            w_tc_borders.append(self.__copy_border__(from_tc_pr, "w:right"))
            to_tc_pr.append(w_tc_borders)

        if len(from_tc_pr.xpath("w:tcMar")) > 0:
            w_tc_mars = OxmlElement("w:tcMar")
            w_tc_mars.append(self.__copy_margin__(from_tc_pr, "w:top"))
            w_tc_mars.append(self.__copy_margin__(from_tc_pr, "w:bottom"))
            w_tc_mars.append(self.__copy_margin__(from_tc_pr, "w:left"))
            w_tc_mars.append(self.__copy_margin__(from_tc_pr, "w:right"))
            to_tc_pr.append(w_tc_mars)


        if len(from_tc_pr.xpath("w:tcW")) > 0:
            w_tc_w = OxmlElement("w:tcW")
            self.__extract_attribute__(from_tc_pr, w_tc_w, "w:tcW", "w:type")
            self.__extract_attribute__(from_tc_pr, w_tc_w, "w:tcW", "w:w")
            to_tc_pr.append(w_tc_w)


class PuppeteerContent(object):

    def __init__(self, puppeteer_extractor, puppeteer_config_name):
        self.puppeteer_extractor = puppeteer_extractor
        self.puppeteer_config_name = puppeteer_config_name

    def apply(self, paragraph, document, initial_run, table_style):
        """
        Insert the placeholder data inside the location defined by the initial_run inside the paragraph.
        The placeholder will be replaced by the information retrieved by Puppeteer, which can be a text, a table or an image.
        :param paragraph: the location of the placeholder
        :param document: the current document, it will be used to retrieve style
        :param initial_run: the run of the placeholder
        :param table_style: a potential style to apply to the current table.
        :return: a new paragraph after the element insertion.
        """
        if isinstance(document, _Cell):
            page_width = document.width
        else:
            first_section = document.sections[0]
            page_width = first_section.page_width - first_section.left_margin - first_section.right_margin
        contents = self.puppeteer_extractor.get_contents(self.puppeteer_config_name)

        # Corner case for model.other_algorithms_search_strategy.image: the first element of each group is a title.
        is_title = self.puppeteer_config_name == "DESIGN_OTHER_ALGORITHMS_SEARCH_STRATEGY_IMAGE" or \
                            self.puppeteer_config_name == "DESIGN_OTHER_ALGORITHMS_SEARCH_STRATEGY_TABLE" or \
                            self.puppeteer_config_name == "ROC_CURVE_MULTICLASS" or \
                            self.puppeteer_config_name == "DENSITY_GRAPH_MULTICLASS" or \
                            self.puppeteer_config_name == "CALIBRATION_CHART_MULTICLASS" or \
                            self.puppeteer_config_name == "BINARY_C_DETAILED_METRICS_TABLE"

        if contents:
            for element_index, content in sorted(contents.items()):
                # Some content are a set of images
                for image_index, extracted_content in sorted(content.items()):
                    if extracted_content["type"] == "txt":
                        paragraph = self.apply_templated_text(extracted_content["data"], paragraph, initial_run, document, is_title)
                        # model.other_algorithms_search_strategy.image is one title, one normal text and one image => next one is not a title
                        is_title = False
                    elif extracted_content["type"] == "png":
                        paragraph = self.apply_chart(extracted_content["data"], paragraph, initial_run, page_width)
                        # model.other_algorithms_search_strategy.image is one title, one normal text and one image => next one is a title
                        is_title = True
                    elif extracted_content["type"] == "json":
                        paragraph = self.apply_templated_json(extracted_content["data"], paragraph, initial_run, document, table_style)
                        # model.other_algorithms_search_strategy.image is one title, one normal text and one image => next one is a title
                        is_title = True
        return paragraph

    def apply_templated_text(self, extracted_text, paragraph, initial_run, document, is_title):
        if not extracted_text:
            return paragraph

        lines = extracted_text.splitlines()
        lines = list(filter(None, lines)) # Removes empty lines

        BEFORE_H1_PLACEHOLDER = "[#@#]"  # This value should match the one used in export-charts.js
        lines = list(map(lambda l: safe_unicode_str(l).replace(BEFORE_H1_PLACEHOLDER, ""), lines))  # Add an empty line before each <h1> to better separate text sections

        first_line = lines.pop(0)

        paragraph_style = paragraph.style
        title_style = paragraph.style
        if is_title:
            # if we are on a title, we want have an empty paragraph and then create another paragraph just for the title
            paragraph = DocumentHandler.insert_paragraph_after(paragraph, style=paragraph_style)
            styles = document.styles

            # use Heading 5 style if this exists.
            paragraph_styles = [
                s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH
            ]
            for style in paragraph_styles:
                if "Heading 5" == style.name:
                    title_style = "Heading 5"
            paragraph = DocumentHandler.insert_paragraph_after(paragraph, style=title_style)
            run = paragraph.add_run("")
        else:
            # First line on the current paragraph
            run = paragraph.add_run("", initial_run.style)


        DocumentHandler.copy_font(initial_run, run)
        run.text += first_line

        # Create a new paragraph for each next line
        for line in lines:
            paragraph = DocumentHandler.insert_paragraph_after(paragraph, style=paragraph_style)
            run = paragraph.add_run("", initial_run.style)
            DocumentHandler.copy_font(initial_run, run)
            run.text += line

        if is_title:
            if title_style != "Heading 5":
                # If Heading 5 does not exists, fake the style of the title: bolder, underlined and bigger.
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.underline = True
                    if paragraph.style.font.size:
                        run.font.size = paragraph.style.font.size + 63500
            # Force new paragraph for next element
            paragraph = DocumentHandler.insert_paragraph_after(paragraph, style=paragraph_style)
        return paragraph

    def apply_chart(self, image, paragraph, initial_run, page_width):
        new_paragraph = DocumentHandler.insert_paragraph_after(paragraph, style=paragraph.style)
        run = new_paragraph.add_run("", initial_run.style)
        DocumentHandler.copy_font(initial_run, run)
        logging.debug("Format for %s is %s", self.puppeteer_config_name, DocxParser.debug_format(paragraph.paragraph_format))
        inline_shape = run.add_picture(BytesIO(image))

        # Because the Puppeteer browser takes high resolution screenshots (deviceScaleFactor=2),
        # we need to divide the image dimensions by 2 in the docx so that it appears at the normal size to the user
        inline_shape.height = int(inline_shape.height / 2)
        inline_shape.width = int(inline_shape.width / 2)

        # If the image is bigger than the page, reduce it
        if inline_shape.width > page_width:
            if inline_shape.width != 0:
                inline_shape.height = inline_shape.height * page_width // inline_shape.width
            inline_shape.width = page_width

        return new_paragraph

    def apply_templated_json(self, extracted_text, paragraph, initial_run, document, table_style):
        if not extracted_text:
            return paragraph

        # The format of the json is:
        # [
        #   ["line1 col1", "line1 col2", "line1 col3 subline1\nline1 col3 subline2"],
        #   ["line2 col1", "line2 col2"]
        # ]
        values = json.loads(extracted_text)
        cells = []
        for line in values:
            cells.append(line)
        table = Table(cells)
        return table.apply(paragraph, document, initial_run, table_style)
