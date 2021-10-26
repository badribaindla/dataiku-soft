from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement, qn
from collections import OrderedDict
import logging

class DocumentHandler(object):
    """
    Extends the docx API to overcome its limitations
    """
    @staticmethod
    def delete_paragraph(paragraph):
        p = paragraph._element
        if p.getparent() is not None:
            p.getparent().remove(p)
        p._p = p._element = None

    @staticmethod
    def delete_table(table):
        tbl = table._element
        if tbl.getparent() is not None:
            tbl.getparent().remove(tbl)
        tbl._tbl = tbl._element = None

    @staticmethod
    def get_paragraph(document, paragraph_pointer):
        for paragraph in document.paragraphs:
            if paragraph._p == paragraph_pointer:
                return paragraph
        return None

    @staticmethod
    def delete_element(element):
        if isinstance(element, Table):
            DocumentHandler.delete_table(element)
        elif isinstance(element, Paragraph):
            DocumentHandler.delete_paragraph(element)
        else:
            logging.warn("Unable to delete element : unsupported type %s", element.__class__.__name__)

    @staticmethod
    def between(document, block):
        logging.debug("Extracting elements between %s and %s", block.start, block.end)
        elements = DocumentHandler.elements(document)
        next = DocumentHandler.next(elements, block.start.paragraph._p)
        extracted = []
        while next is not None and next != block.end.paragraph._p:
            extracted.append(next)
            next = DocumentHandler.next(elements, next)
        otw = DocumentHandler.ooxml_to_wrapper(document)
        wrappers = []
        for e in extracted:
            if e in otw:
                wrappers.append(otw[e])
        logging.debug("Extracted elements are %s", wrappers)
        return wrappers

    @staticmethod
    def next(elements, key):
        if hasattr(elements, "_OrderedDict__map"):
            # Python2
            if key in elements._OrderedDict__map:
                link_prev, link_next, key = elements._OrderedDict__map[key]
                return link_next[2]
            else:
                return None
        else:
            # Python3
            if key in elements:
                keys = list(elements)
                key_index = keys.index(key)
                if key_index < len(keys) - 1:
                    return keys[key_index + 1]
            return None

    @staticmethod
    def ooxml_to_wrapper(document):
        elt_ids = {}
        for p in document.paragraphs:
            elt_ids[p._p] = p
        for tbl in document.tables:
            elt_ids[tbl._tbl] = tbl
        return elt_ids

    @staticmethod
    def elements(document):
        """
        Returns the dict of the elements.
        This dict is ordered. We use a dict to easily access next and previous elements of an item.
        We must use it because a `docx.Document` does not expose a combined list of paragraphs and tables but maintains two separated lists.
        Internally, he maintains a ordered list of ooxml element mixing paragraph and tables
        Or when dealing with block placeholders (aka {{placeholder}}TABLE{{placeholder}}), 
        we need to find the proper table between the two placeholders.
        """
        elts = OrderedDict()
        if isinstance(document, _Cell):
            for elt in document._element:
                try:
                    elts[elt] = elt
                except:
                    logging.error("%s not found", elt)
        else:
            for elt in document._body._element:
                try:
                    elts[elt] = elt
                except:
                    logging.error("%s not found", elt)
        return elts

    @staticmethod
    def insert_paragraph_after(paragraph, text=None, style=None):
        """Insert a new paragraph after the given paragraph."""
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        new_para = Paragraph(new_p, paragraph._parent)
        if text:
            new_para.add_run(text)
        if style is not None:
            new_para.style = style
        DocumentHandler.copy_paragraph_format(paragraph.paragraph_format, new_para.paragraph_format)
        return new_para

    @staticmethod
    def copy_paragraph_format(src, dest):
        """
        Copy paragraph format attributes to an other one.
        Warning: it is impossible to copy tab_stops as the docx paragraph_format api doesn't offers any setter
        """
        dest.alignment = src.alignment
        dest.first_line_indent = src.first_line_indent
        dest.keep_together = src.keep_together
        dest.keep_with_next = src.keep_with_next
        dest.left_indent = src.left_indent
        dest.line_spacing = src.line_spacing
        dest.line_spacing_rule = src.line_spacing_rule
        dest.page_break_before = src.page_break_before
        dest.right_indent = src.right_indent
        dest.space_after = src.space_after
        dest.space_before = src.space_before
        dest.widow_control = src.widow_control

    @staticmethod
    def copy_font(from_run, to_run):
        to_run.font.name = from_run.font.name
        to_run.font.size = from_run.font.size
        to_run.font.italic = from_run.font.italic
        to_run.font.cs_italic = from_run.font.cs_italic
        to_run.font.bold = from_run.font.bold
        to_run.font.cs_bold = from_run.font.cs_bold
        to_run.font.underline = from_run.font.underline
        to_run.font.strike = from_run.font.strike
        if from_run.font.color:
            if from_run.font.color.rgb:
                to_run.font.color.rgb = from_run.font.color.rgb
            if from_run.font.color.theme_color:
                to_run.font.color.theme_color = from_run.font.color.theme_color
                w_theme_shade = from_run.font.color._color.xpath("@w:themeShade")
                if len(w_theme_shade) > 0:
                    to_run.font.color._color.set(qn("w:themeShade"), w_theme_shade[0])
                w_theme_tint = from_run.font.color._color.xpath("@w:themeTint")
                if len(w_theme_tint) > 0:
                    to_run.font.color._color.set(qn("w:themeTint"), w_theme_tint[0])

