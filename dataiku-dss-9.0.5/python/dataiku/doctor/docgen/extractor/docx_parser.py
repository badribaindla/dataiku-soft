from dataiku.doctor.docgen.extractor.placeholder_parser import PlaceholderParser
from docx.styles.style import _ParagraphStyle
from docx.text.run import Run
from dataiku.doctor.docgen.common.placeholder import Placeholder, TablePlaceholder, BlockPlaceholder, \
    HeaderPlaceholder, FooterPlaceholder
import logging


class DocxParser(object):

    @staticmethod
    def get_safe_runs(p):
        """
        Sequence of |Run| instances corresponding to the <w:r> elements in the p paragraph.
        The difference with p.runs is that this method also includes the runs from accepted changes
        (meaning runs inserted when the "Track Changes" mode is enabled on Word)
        """
        return [Run(r, p) for r in p._p.xpath(".//w:r[not(ancestor::w:del)]")]

    @staticmethod
    def parse_text(document):
        """
        Find placeholders that are located in the plain text.
        :returns: a list of placeholders
        :rtype: [Placeholder]
        """
        return [
            Placeholder(name, is_conditional, is_closing, para, start, end)
            for para in document.paragraphs
            for name, is_conditional, is_closing, start, end in PlaceholderParser.parse(DocxParser.get_safe_runs(para))
        ]

    def parse_table(self, document):
        """
        Find placeholders that are located inside table.
        :returns: a list of placeholders
        :rtype: [TablePlaceholder]
        """
        placeholder_tables = []
        for table in document.tables:
            for row_id, row in enumerate(table.rows):
                for column_id, cell in enumerate(row.cells):
                    placeholders = self.parse_text(cell)
                    if len(placeholders) > 0:
                        placeholder_tables.append(TablePlaceholder(table, row_id, column_id, placeholders))
        return placeholder_tables

    def parse_headers(self, document):
        """
        Find placeholders that are located inside headers.
        :returns: a list of placeholders
        :rtype: [HeaderPlaceholder]
        """
        return [
            HeaderPlaceholder(name, is_conditional, is_closing, para, start, end, section_number)
            # For each section
            for section_number in range(len(document.sections))
            # get the paragraphs of the header
            for para in document.sections[section_number].header.paragraphs
            # and extract the placeholders
            for name, is_conditional, is_closing, start, end in PlaceholderParser.parse(DocxParser.get_safe_runs(para))
        ]

    def parse_footers(self, document):
        """
        Find placeholders that are located inside footers.
        :returns: a list of placeholders
        :rtype: [FooterPlaceholder]
        """
        return [
            FooterPlaceholder(name, is_conditional, is_closing, para, start, end, section_number)
            # For each section
            for section_number in range(len(document.sections))
            # get the paragraphs of the footer
            for para in document.sections[section_number].footer.paragraphs
            # and extract the placeholders
            for name, is_conditional, is_closing, start, end in PlaceholderParser.parse(DocxParser.get_safe_runs(para))
        ]

    @staticmethod
    def get_end_conditionals_placeholders(placeholder_name, iterator, placeholders):
        """
        find the endif placeholder associated to the placeholder named "placeholder_name", starting from the iterator
        :param [Placeholder] placeholders: list of placeholder
        :param placeholder_name: placeholder name
        :param iterator: starting point
        :return: if we find a matching placeholder, return it, otherwise return None.
        """
        # On some really unwanted case, we can have multiple conditional placeholder using the same placeholder variable
        # We need to match the "if" with the correct "endif".
        recursivity_counter = 0
        for placeholder in placeholders[iterator:]:
            if placeholder.is_conditional and placeholder.is_closing and \
                    placeholder_name == placeholder.extract_name():
                if recursivity_counter > 0:
                    # We reach an endif that is associated to an inner conditional placeholder
                    recursivity_counter -= 1
                else:
                    return placeholder
            elif placeholder.is_conditional and not placeholder.is_closing and \
                placeholder_name == placeholder.extract_name():
                # we have an inner conditional placeholder. We try to see if the two placeholders names match.
                # if so, we increment the recursivity_counter to not take the first endif.
                recursivity_counter += 1
        return None

    @staticmethod
    def extract_conditionals_placeholders(placeholders):
        """
        Extract from the placeholders list the conditionals placeholders

        ie:
        {if myplaceholder == myvalue}
        a text to display.
        {endif myplaceholder}

        :param placeholders: List of extracted placeholders.
        :return: [:class:`BlockPlaceholder`]
        """
        conditional_placeholders = []
        i = 0
        while i < len(placeholders):
            if placeholders[i].is_conditional and not placeholders[i].is_closing:
                real_name = placeholders[i].extract_name()
                logging.debug("name: '%s' => '%s'", placeholders[i].tagname, real_name)
                closing_placeholder = DocxParser.get_end_conditionals_placeholders(real_name, i + 1, placeholders)
                if closing_placeholder is not None:
                    conditional_placeholders.append(BlockPlaceholder(placeholders[i], closing_placeholder))
                else:
                    logging.error("No match for conditional placeholder : '%s'.", placeholders[i].tagname)
            # else: non conditional or closing placeholder. We do not care
            i += 1
        return conditional_placeholders

    @staticmethod
    def extract_blocks(placeholders):
        """
        Given a list a placeholders, returns the blocks from that list and the original list without the blocks placeholders
        A block is defined as two placeholders with the same name following each others.
        The second placeholder must start with a /

        ie: {mytable} table description with style{/mytable}

        :param [Placeholder] placeholders: list of placeholder extracted from the docx file.
        :rtype: ([:class:`BlockPlaceholder`], [single_placeholders])
        """
        logging.debug("Placeholders before extraction :%s", placeholders)
        single_placeholders = []
        blocks = []
        i = 0
        while i < len(placeholders):
            # remove conditional placeholders
            if not placeholders[i].is_conditional:
                if not placeholders[i].is_closing and i < len(placeholders) - 1 and \
                        placeholders[i + 1].is_closing and placeholders[i].tagname == placeholders[i + 1].tagname:
                    # check if we are on a block placeholder (aka next paragraph is a closing one)
                    blocks.append(BlockPlaceholder(placeholders[i], placeholders[i + 1]))
                    i += 1
                elif not placeholders[i].is_closing:
                    single_placeholders.append(placeholders[i])
                # else: closing placeholder it was already processed or is an error and should be left as is
            i += 1

        logging.debug("Blocks found : %s", blocks)
        logging.debug("Placeholders left : %s", single_placeholders)
        return blocks, single_placeholders

    def debug(self, document):
        for i, p in enumerate(document.paragraphs):
            logging.debug("Paragraph: %d %s format: %s style: %s", i, hex(id(p._element)),
                          self.debug_format(p.paragraph_format), self.debug_c_style(p.style))
            for j, r in enumerate(document.paragraphs[i].runs):
                logging.debug("\tRun [%d] : %s. style:%s, bold=%s, italic=%s, font=%s",
                              j, r.text, self.debug_c_style(r.style), r.bold, r.italic, self.debug_font(r.font))
            logging.debug("Text %s", p.text)

        # too spammy
        # logging.debug("Availables styles")

        # for style in document.styles:
        #    logging.debug("%s:%s format %s", style.name, style.element,
        #                  self.debug_format(style.paragraph_format) if isinstance(style, _ParagraphStyle) else "")
        for table in document.tables:
            logging.debug("Table: %s %s parent: %s", table, table._element, table._parent)
            for row in table.rows:
                logging.debug("====")
                for cell in row.cells:
                    logging.debug("cell Text: %s ", cell.text)
        logging.debug("Body:%s ", document._body._element)
        for i in document._body._element:
            logging.debug(i)

    def debug_c_style(self, c_style):
        if c_style is None:
            return "None"
        return "{" + c_style.name + "font=" + self.debug_font(c_style.font) + " base=" \
               + self.debug_c_style(c_style.base_style) + "}"

    @staticmethod
    def debug_format(format):
        return str(format) + " left_indent: " + str(format.left_indent) + " " + format.left_indent.__class__.__name__

    @staticmethod
    def debug_font(font):
        items = {
            "bold": font.bold,
            "color": font.color,
            "italic": font.italic,
            "name": font.name,
            "cs_bold": font.cs_bold,
            "cs_italic": font.cs_italic,
            "size": font.size}
        return ", ".join([attr + "=" + str(value) for attr, value in items.items()])
